import os
import uuid
import asyncio
from datetime import datetime
from pathlib import Path
from typing import Optional, Tuple
import aiofiles
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.utils import ImageReader
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
import openpyxl
from openpyxl import Workbook
import PyPDF2
import pdfplumber
import tempfile
import subprocess
import platform

# Create conversion directory if it doesn't exist
CONVERSION_DIR = Path("/tmp/conversions")
CONVERSION_DIR.mkdir(exist_ok=True)

class ConversionService:
    """Service for handling file conversions between different formats"""
    
    @staticmethod
    def generate_unique_filename(original_name: str, new_extension: str) -> str:
        """Generate a unique filename with the new extension"""
        base_name = Path(original_name).stem
        unique_id = str(uuid.uuid4())[:8]
        return f"{base_name}_{unique_id}.{new_extension}"
    
    @staticmethod
    async def save_uploaded_file(file, filename: str) -> str:
        """Save uploaded file to conversion directory"""
        file_path = CONVERSION_DIR / filename
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        return str(file_path)
    
    @staticmethod
    def cleanup_file(file_path: str):
        """Remove temporary file"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception:
            pass  # Ignore cleanup errors
    
    # Word to PDF Conversion
    @staticmethod
    async def word_to_pdf(input_path: str, output_path: str) -> bool:
        """Convert Word document to PDF"""
        try:
            # Try LibreOffice first for all systems
            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', str(Path(output_path).parent),
                input_path
            ], capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0:
                # LibreOffice creates PDF with same base name as input
                input_base = Path(input_path).stem
                generated_pdf = Path(output_path).parent / f"{input_base}.pdf"
                
                # Move to expected output path if different
                if generated_pdf.exists() and str(generated_pdf) != output_path:
                    os.rename(str(generated_pdf), output_path)
            else:
                # Fallback: Manual conversion using python-docx and reportlab
                await ConversionService._manual_word_to_pdf(input_path, output_path)
            
            return os.path.exists(output_path)
        except Exception as e:
            print(f"Word to PDF conversion error: {e}")
            return False
    
    @staticmethod
    async def _manual_word_to_pdf(input_path: str, output_path: str):
        """Manual Word to PDF conversion using python-docx + reportlab"""
        doc = Document(input_path)
        
        # Create PDF
        pdf_doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Process paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Determine style based on paragraph format
                if para.style.name.startswith('Heading'):
                    style = styles['Heading1']
                else:
                    style = styles['Normal']
                
                p = Paragraph(para.text, style)
                story.append(p)
                story.append(Spacer(1, 6))
        
        pdf_doc.build(story)
    
    # PDF to Word Conversion
    @staticmethod
    async def pdf_to_word(input_path: str, output_path: str) -> bool:
        """Convert PDF to Word document"""
        try:
            # Use manual conversion method (fallback approach)
            await ConversionService._manual_pdf_to_word(input_path, output_path)
            return os.path.exists(output_path)
        except Exception as e:
            print(f"PDF to Word conversion error: {e}")
            return False
    
    @staticmethod
    async def _manual_pdf_to_word(input_path: str, output_path: str):
        """Manual PDF to Word conversion using text extraction"""
        # Extract text from PDF
        text_content = []
        
        with open(input_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
        
        # Create Word document
        doc = Document()
        
        for i, page_text in enumerate(text_content):
            if i > 0:
                doc.add_page_break()
            
            # Add paragraphs (split by double newlines)
            paragraphs = page_text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
        
        doc.save(output_path)
    
    # Text to PDF Conversion
    @staticmethod
    async def text_to_pdf(input_path: str, output_path: str) -> bool:
        """Convert text file to PDF"""
        try:
            async with aiofiles.open(input_path, 'r', encoding='utf-8') as f:
                text_content = await f.read()
            
            # Create PDF
            pdf_doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Split text into paragraphs
            paragraphs = text_content.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Replace single newlines with <br/> tags for proper formatting
                    formatted_text = para_text.replace('\n', '<br/>')
                    p = Paragraph(formatted_text, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))
            
            if not story:
                # If no paragraphs, add the entire text as one block
                p = Paragraph(text_content.replace('\n', '<br/>'), styles['Normal'])
                story.append(p)
            
            pdf_doc.build(story)
            return True
            
        except Exception as e:
            print(f"Text to PDF conversion error: {e}")
            return False
    
    # Image to PDF Conversion
    @staticmethod
    async def image_to_pdf(input_paths: list, output_path: str) -> bool:
        """Convert image(s) to PDF"""
        try:
            if len(input_paths) == 1:
                # Single image
                image = Image.open(input_paths[0])
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                image.save(output_path, "PDF")
            else:
                # Multiple images
                images = []
                for img_path in input_paths:
                    img = Image.open(img_path)
                    if img.mode != 'RGB':
                        img = img.convert('RGB')
                    images.append(img)
                
                # Save first image with the rest appended
                images[0].save(output_path, "PDF", save_all=True, append_images=images[1:])
            
            return True
            
        except Exception as e:
            print(f"Image to PDF conversion error: {e}")
            return False
    
    # Excel to PDF Conversion
    @staticmethod
    async def excel_to_pdf(input_path: str, output_path: str) -> bool:
        """Convert Excel file to PDF"""
        try:
            # Try LibreOffice first (most reliable for Excel to PDF)
            if platform.system() != "Windows":
                result = subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', str(Path(output_path).parent),
                    input_path
                ], capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0:
                    # LibreOffice creates PDF with same base name as input
                    input_base = Path(input_path).stem
                    generated_pdf = Path(output_path).parent / f"{input_base}.pdf"
                    
                    # Move to expected output path if different
                    if generated_pdf.exists() and str(generated_pdf) != output_path:
                        os.rename(str(generated_pdf), output_path)
                    
                    return os.path.exists(output_path)
            
            # Fallback: Manual conversion using openpyxl + reportlab
            await ConversionService._manual_excel_to_pdf(input_path, output_path)
            return True
            
        except Exception as e:
            print(f"Excel to PDF conversion error: {e}")
            return False
    
    @staticmethod
    async def _manual_excel_to_pdf(input_path: str, output_path: str):
        """Manual Excel to PDF conversion"""
        # Read Excel file
        workbook = openpyxl.load_workbook(input_path, data_only=True)
        
        # Create PDF
        c = canvas.Canvas(output_path, pagesize=A4)
        width, height = A4
        
        y_position = height - 50
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # Add sheet title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, y_position, f"Sheet: {sheet_name}")
            y_position -= 30
            
            # Add table headers and data
            c.setFont("Helvetica", 10)
            
            for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
                if y_position < 50:  # Start new page
                    c.showPage()
                    y_position = height - 50
                
                x_position = 50
                for cell_value in row:
                    if cell_value is not None:
                        cell_text = str(cell_value)[:20]  # Truncate long text
                        c.drawString(x_position, y_position, cell_text)
                    x_position += 100
                
                y_position -= 15
                
                if row_idx > 100:  # Limit rows to prevent huge files
                    break
            
            # Add page break between sheets
            if sheet_name != workbook.sheetnames[-1]:
                c.showPage()
                y_position = height - 50
        
        c.save()
    
    # File validation
    @staticmethod
    def validate_file(file_path: str, allowed_extensions: list) -> Tuple[bool, str]:
        """Validate file type and size"""
        try:
            file_ext = Path(file_path).suffix.lower().lstrip('.')
            
            # Check extension
            if file_ext not in allowed_extensions:
                return False, f"Invalid file type. Allowed: {', '.join(allowed_extensions)}"
            
            # Check file size (50MB limit)
            file_size = os.path.getsize(file_path)
            max_size = 50 * 1024 * 1024  # 50MB
            if file_size > max_size:
                return False, "File size exceeds 50MB limit"
            
            # Check if file exists and is readable
            if not os.path.exists(file_path):
                return False, "File not found"
            
            return True, "Valid file"
            
        except Exception as e:
            return False, f"File validation error: {str(e)}"